# -*- coding: utf-8 -*-
"""生成《肠道花园 Gut Garden》核心代码整理 docx（约 10 页）。
代码片段直接从源码文件按行截取，保证与真实代码一致。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EAST = '微软雅黑'

def read_lines(rel_path, start, end):
    """读取源码文件第 start..end 行（含两端），返回去掉首尾空行后的字符串。"""
    full = os.path.join(ROOT, rel_path)
    with open(full, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    block = lines[start - 1:end]
    while block and not block[0].strip():
        block = block[1:]
    while block and not block[-1].strip():
        block = block[:-1]
    return '\n'.join(block)

def east_asia(run, font=EAST):
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font)

def set_style_font(style, size, bold=False, color=None):
    style.font.name = 'Calibri'
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    rPr = style.element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), EAST)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.3); s.right_margin = Cm(2.3)

set_style_font(doc.styles['Normal'], 11)
set_style_font(doc.styles['Heading 1'], 15, bold=True, color=(0x2E, 0x5B, 0x2A))
set_style_font(doc.styles['Heading 2'], 12.5, bold=True, color=(0x2E, 0x5B, 0x2A))

def p(text, size=11, bold=False, color=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    east_asia(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.3
    return para

def bullet(text, size=10.5):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    east_asia(run)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(3)
    return para

def code_block(title, text):
    if title:
        t = doc.add_paragraph()
        r = t.add_run(title)
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        east_asia(r, '宋体')
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        t.paragraph_format.space_before = Pt(4)
        t.paragraph_format.space_after = Pt(2)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.right_indent = Cm(0.4)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    east_asia(run, '宋体')
    run.font.color.rgb = RGBColor(0x1F, 0x3D, 0x2E)
    rPr = run._element.get_or_add_rPr()
    shd = rPr.makeelement(qn('w:shd'), {qn('w:fill'): 'F4F6F0'})
    rPr.append(shd)
    pPr = para._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn('w:pBdr'), {})
    for side in ('top', 'left', 'bottom', 'right'):
        el = pPr.makeelement(qn('w:' + side), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '4', qn('w:color'): 'B7C4AE'})
        pbdr.append(el)
    pPr.append(pbdr)
    return para

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        r = hdr[i].paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        east_asia(r)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(9)
            east_asia(r)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================== 封面标题 ==============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('肠道花园 Gut Garden\n核心代码整理')
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(0x2E, 0x5B, 0x2A)
east_asia(r)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('—— 把看不见的肠道微生态，做成孩子能看、能玩、能懂、能养成习惯的互动科普\n前端 React + 后端 Fastify · 全栈 TypeScript · 内置 PGlite 数据库 · AI 全链路可降级')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
east_asia(r)
doc.add_paragraph()

# ============================== 总览 ==============================
doc.add_heading('〇、核心代码总览', level=1)
p('本项目是「一个 IP、三类赛道」的互动科普装置，核心代码按「内容数据化、AI 可降级、行为即数据」三条原则组织。下表列出本次整理的代码文件与其承担的角色：')
table(
    ['模块', '文件', '承担的核心职责'],
    [
        ['AI 导游', 'server/src/modules/ai/*', '菌小园科普问答：知识底座注入 system prompt，流式 SSE 输出，未配置 Key/上游失败自动降级本地 FAQ，零幻觉'],
        ['便便垂直 AI', 'server/src/modules/stool/*', '便便照片/形态分析：结构化 JSON 输出约束，识别 Bristol 1-7 型并生成专属饮食任务'],
        ['每日打卡', 'server/src/modules/checkin/checkin.service.ts', '5 个健康任务打卡、连续打卡计算、补卡、同步日历'],
        ['徽章引擎', 'server/src/modules/badges/engine.ts', '事件驱动发奖：收集行为统计 → 规则匹配 → 铜银金升级，幂等防重'],
        ['花园成长', 'server/src/modules/garden/stage.service.ts', '6 阶段成长规则数据化，打卡/投喂/徽章三项指标驱动'],
        ['花园投喂', 'web/src/hooks/useFeedLogic.ts', '前端拖拽食物 → 花园实时反馈；注册用户写后端、游客本地状态'],
        ['验证码登录', 'server/src/modules/auth/*', '手机号+验证码（模拟短信），任意号码自动注册，60s 限频 / 5min 有效'],
    ],
    widths=[2.6, 6.0, 7.6],
)

# ============================== 一、AI 导游 ==============================
doc.add_heading('一、AI 导游「菌小园」：知识底座注入 + 流式输出 + 兜底降级', level=1)
p('科普问答的底线是「科学准确」。我们不允许大模型自由发挥，而是把应用内已验证的知识（知识模块 / 便便分型 / 花园阶段 / 常见问答）动态组装成「知识底座」注入 system prompt，AI 只能依据底座回答；模型不可用或超时时，自动降级为本地 FAQ，体验永不中断。')

doc.add_heading('1.1 知识底座：从已验证数据文件动态生成，杜绝幻觉', level=2)
p('知识全部来自数据文件，改一条内容不用动代码，换主题只需换数据：')
code_block('server/src/modules/ai/knowledge-base.ts — buildKnowledgeText()',
    read_lines('server/src/modules/ai/knowledge-base.ts', 6, 43))

doc.add_heading('1.2 组装消息：系统提示 + 页面上下文 + 最近 6 轮历史', level=2)
code_block('server/src/modules/ai/ai.routes.ts — buildMessages()',
    read_lines('server/src/modules/ai/ai.routes.ts', 11, 23))

doc.add_heading('1.3 流式输出与三级降级链', level=2)
p('有 Key 走 OpenAI 兼容流式接口；上游失败、无输出、没 Key 都落到本地 FAQ，保证「菌小园」永远有回应：')
code_block('server/src/modules/ai/ai-stream.ts — streamAiChunks()',
    read_lines('server/src/modules/ai/ai-stream.ts', 50, 69))

# ============================== 二、便便垂直 AI ==============================
doc.add_heading('二、便便观察「垂直 AI」：结构化输出约束 + 预设降级', level=1)
p('便便建议做成「垂直专业 AI」不需要微调模型，我们用「强约束 system prompt + 结构化 JSON 输出 + 预设兜底」三件套实现。模型只负责按知识底座「查找 + 翻译」，且输出被硬性格式约束，天然可解析、可回退。')

doc.add_heading('2.1 垂直角色 system prompt（角色 / 知识底座 / 输出格式 / 硬性规则）', level=2)
code_block('server/src/modules/stool/stool-ai.ts — STOOL_AI_SYSTEM_PROMPT()',
    read_lines('server/src/modules/stool/stool-ai.ts', 13, 45))

doc.add_heading('2.2 生成建议：15s 超时，输出不可解析一律回退预设', level=2)
code_block('server/src/modules/stool/stool-ai.ts — generateStoolSuggestion()',
    read_lines('server/src/modules/stool/stool-ai.ts', 80, 115))

doc.add_heading('2.3 照片分析客户端：外部 API 与本地规则双通道', level=2)
code_block('server/src/modules/stool/stool-analysis.client.ts — analyzeStoolPhoto() + mockAnalyze()',
    read_lines('server/src/modules/stool/stool-analysis.client.ts', 14, 29) + '\n\n' +
    read_lines('server/src/modules/stool/stool-analysis.client.ts', 69, 75))

# ============================== 三、每日打卡 ==============================
doc.add_heading('三、每日打卡：数据驱动任务 + 连续打卡计算', level=1)
p('打卡系统把「任务定义」做成数据表，5 个健康任务（探索花园 / 健康饮食 / 优质睡眠 / 补充水分 / 活力运动）对应同一张记录表的 5 列，任务是否完成由代码统一判定，便于日历同步与徽章联动。')

doc.add_heading('3.1 任务定义：数据驱动', level=2)
code_block('server/src/modules/checkin/checkin.service.ts — TASK_DEFS',
    read_lines('server/src/modules/checkin/checkin.service.ts', 9, 17))

doc.add_heading('3.2 连续打卡计算（当前连续 & 历史最长）', level=2)
code_block('server/src/modules/checkin/checkin.service.ts — computeStreaks()',
    read_lines('server/src/modules/checkin/checkin.service.ts', 60, 96))

doc.add_heading('3.3 确认任务：事务加经验 + 联动徽章', level=2)
code_block('server/src/modules/checkin/checkin.service.ts — confirmTask()',
    read_lines('server/src/modules/checkin/checkin.service.ts', 178, 196))

# ============================== 四、徽章引擎 ==============================
doc.add_heading('四、徽章引擎：事件驱动 + 幂等发放', level=1)
p('徽章规则存在数据库（badge_defs 的 condition_rule / silver_rule / gold_rule 三档），引擎统一收集行为统计，逐条规则判断，按铜→银→金依次升级。通过 event_id 唯一索引保证幂等：同一枚徽章重复触发不会重复发奖。')

doc.add_heading('4.1 规则匹配：把「规则即数据」翻译成判定', level=2)
code_block('server/src/modules/badges/engine.ts — checkRule()',
    read_lines('server/src/modules/badges/engine.ts', 154, 190))

doc.add_heading('4.2 事件驱动评估：铜银金三档 + 幂等落库', level=2)
code_block('server/src/modules/badges/engine.ts — evaluate()',
    read_lines('server/src/modules/badges/engine.ts', 199, 238))

# ============================== 五、花园成长 ==============================
doc.add_heading('五、花园成长：阶段规则数据化', level=1)
p('花园 6 个成长阶段（种子 → 幼苗 → 成长 → 丰收 → 大师 → 终极）用一条规则表描述，换数值只改数据不改逻辑：')
code_block('server/src/modules/garden/stage.service.ts — STAGE_REQS + evaluateStage()',
    read_lines('server/src/modules/garden/stage.service.ts', 17, 24) + '\n\n' +
    read_lines('server/src/modules/garden/stage.service.ts', 47, 58))

# ============================== 六、花园投喂 ==============================
doc.add_heading('六、前端花园投喂：一次拖拽 = 一次行为反馈闭环', level=1)
p('食物被拖进花园后，根据「食物 → 效果」映射表即时反馈：注册用户写后端（记行为、加经验、触发花园状态变化），游客模式直接改本地状态模拟，保证未登录也能玩：')
code_block('web/src/hooks/useFeedLogic.ts — handleDrop()',
    read_lines('web/src/hooks/useFeedLogic.ts', 29, 82))

# ============================== 七、验证码登录 ==============================
doc.add_heading('七、验证码登录（模拟短信）：限频 + 过期 + 自动注册', level=1)
p('项目用「手机号 + 验证码」登录，不发真实短信。验证码 60 秒内同一号码只能发一次、5 分钟有效，校验通过后按号码自动注册家长账号，方便演示与测试：')
code_block('server/src/modules/auth/codeStore.ts — 验证码存储与校验',
    read_lines('server/src/modules/auth/codeStore.ts', 1, 24))
code_block('server/src/modules/auth/auth.service.ts — loginWithCode()',
    read_lines('server/src/modules/auth/auth.service.ts', 10, 24))

# ============================== 结尾 ==============================
doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run('注：以上代码片段均从源码文件按行截取，保证与当前仓库一致。完整代码见 https://github.com/Zeadeinsung/gut-garden —— 启动与密钥配置见仓库 README「快速部署与启动」「如何配置密钥」章节。')
r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
east_asia(r)

OUT = os.path.join(ROOT, 'docs', '比赛方案_核心代码整理.docx')
doc.save(OUT)
print('已生成:', OUT)
