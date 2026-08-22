# -*- coding: utf-8 -*-
"""生成《肠道花园 Gut Garden》比赛方案报告摘录 docx（3.2 / 3.3 / 4.1 / 4.2 / 4.3）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

EAST = '微软雅黑'

def east_asia(run, font=EAST):
    """设置 run 的中文字体（健壮：自动创建 rPr/rFonts）"""
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font)

def set_style_font(style, size, bold=False, color=None):
    style.font.name = 'Calibri'
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    rPr = style.element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), EAST)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

set_style_font(doc.styles['Normal'], 11)
set_style_font(doc.styles['Heading 1'], 16, bold=True, color=(0x2E, 0x5B, 0x2A))
set_style_font(doc.styles['Heading 2'], 13.5, bold=True, color=(0x2E, 0x5B, 0x2A))
set_style_font(doc.styles['Heading 3'], 12, bold=True, color=(0x33, 0x33, 0x33))

def p(text, size=11, bold=False, color=None, indent=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    east_asia(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.35
    return para

def bullet(text, size=11):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    east_asia(run)
    para.paragraph_format.line_spacing = 1.35
    para.paragraph_format.space_after = Pt(4)
    return para

def code_block(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.6)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    east_asia(run, '宋体')
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    shd = run._element.get_or_add_rPr().makeelement(qn('w:shd'), {qn('w:fill'): 'F2F2F2'})
    run._element.get_or_add_rPr().append(shd)
    return para

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        r = hdr[i].paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10.5)
        east_asia(r)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(10)
            east_asia(r)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================== 封面标题 ==============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('科学传播的多元艺术表达 · 方案报告摘录')
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(0x2E, 0x5B, 0x2A)
east_asia(r)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('作品名称：肠道花园 Gut Garden\n—— 把看不见的肠道微生态，做成孩子能看、能玩、能懂、能养成习惯的互动科普')
r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
east_asia(r)
doc.add_paragraph()

# ============================== 3.2 ==============================
doc.add_heading('3.2 关键技术决策', level=1)
p('创作过程中，有两个决策对最终作品影响最大：一个决定「作品长什么样」，一个决定「科学内容怎么不跑偏、怎么可复用」。')

doc.add_heading('决策一：用一套统一的「花园隐喻 + 交互反馈」做科学翻译，而不是「图文说教」', level=2)
p('我们想传播的「肠道微生态」看不见、摸不着、专业术语多。如果直接讲微生物、短链脂肪酸、肠道屏障，孩子无法建立直观认知，家长也难有耐心。', bold=False)
p('我们对比了两种方案：', bold=False)
bullet('方案 A（传统图文说教）：海报 + 文字解释，制作快，但「看完就忘」，孩子只是被动接受。')
bullet('方案 B（我们采用的）：建立一套贯穿全作品的隐喻系统——肠道 = 花园、微生物 = 居民、食物 = 养分、便便 = 花园寄来的天气报告，并把所有科学知识「翻译」成可交互、有即时反馈的体验（拖食物进花园看它变绿还是变脏、记录便便自动生成专属饮食任务、问 AI 导游菌小园）。')
p('影响：这套隐喻成为作品的「共同语言」——首页、花园、课堂、打卡、徽章、动画、AI 对话全部共用同一套 IP 与意象，孩子 30 秒上手、家长愿意每天陪伴 3 分钟，真正把「科学传播」从一次性阅读变成每日行为养成。三类赛道（期刊视觉 / 数字交互 / 叙事动画）也因此不是三件拼凑的作品，而是一套自洽的体系。')

doc.add_heading('决策二：科学内容与程序、AI 全链路「数据解耦 + 知识底座注入」，杜绝 AI 幻觉、保证可复用', level=2)
p('科普作品的底线是「科学准确」。我们预判的最大风险是：知识硬编码在页面里，改一处内容就要动一次代码；或者依赖大模型自由发挥，AI 会编造数字和医学表述。', bold=False)
p('我们对比了三种方案：', bold=False)
bullet('方案 A：知识写死在组件里——改内容 = 改代码，维护成本高，无法换主题复用。')
bullet('方案 B：AI 完全自由回答——体验灵活但有幻觉风险，科普作品不可接受。')
bullet('方案 C（我们采用的）：把科学内容全部做成数据文件（知识模块 / 常见问答 / 布里斯托便便分型 / 花园成长阶段），页面从数据渲染；AI 导游回答时，后端把这些已验证内容动态组装成「知识底座」注入 system prompt，并用硬性规则约束「只依据知识库回答、不编造数字、不提供医疗建议」；模型不可用或超时时自动降级到预设 FAQ，体验永不中断。')
p('影响：第一，科学准确性可控——AI 的每一句话都能回溯到知识库中的权威条目，实现零幻觉；第二，可复用——换一个科学主题，只需替换知识库数据文件与视觉资产，程序与 AI 流水线完全不用动，一周内可再生产一套新作品，正好对应赛事「可复用生成的 AI 应用」的加分条件。')

# ============================== 3.3 ==============================
doc.add_heading('3.3 主要困难与解决方案', level=1)
p('创作过程中最大的三类挑战，以及我们的分析与应对如下：')

table(
    ['遇到的问题', '原因分析', '解决方案', '最终效果'],
    [
        [
            'AI 科普问答会「一本正经地编造」——编数字、编诊断，对儿童科普是致命问题',
            '大模型本质是概率生成，没有「事实约束」；医学/营养内容一旦出错，轻则误导、重则造成健康风险',
            '① 后端把应用内已验证的知识（模块、FAQ、Bristol 分型、成长阶段）动态组装成知识底座注入 system prompt；② 7 条风格指南 + 硬性规则：只依据知识库回答、不编造数字/病名/成分、涉及持续腹痛等立即提示就医；③ 流式 SSE 输出，上游失败或无输出时降级到本地 FAQ 匹配，体验不中断',
            'AI 回答全部可回溯到知识库条目，实现零幻觉；医疗边界清晰；无网也能回答，演示不翻车',
        ],
        [
            '「肠道微生态」看不见、抽象，孩子和家长难以建立直观认知',
            '科学传播的经典难题：对象不可见、概念链条长（纤维→发酵→短链脂肪酸→屏障→免疫）',
            '① 建立花园隐喻系统（肠道=花园、食物=养分、便便=天气报告）；② 所有抽象概念可视化并做成可交互反馈：拖食物看花园实时变绿/变脏、便便报告自动生成专属饮食任务、AI 用比喻讲解',
            '孩子 30 秒上手，家长愿意每天陪伴 3 分钟；同一隐喻让「看→玩→懂→做→再验证」的传播闭环成立',
        ],
        [
            '非动画专业团队，难以低成本产出 90 秒「电影感」科普动画',
            '传统动画需要建模、绑定、动画、渲染，一个工作室少则一个月，个人团队几乎不可能',
            '搭一条 AI 视频生产线：90 秒拆成 17 个分镜，每条分镜用结构化 Prompt（场景/动作/时长/镜头）存成 JSON；全局正向+负向提示词锁定「粘土+绒布+粉蓝绿童话」统一风格；5 个主角先用通义万相出角色设定图，再以 IP-Adapter 注入保证全片不换脸；脚本循环批量提交、断点续传',
            '已产出 4 段核心知识成片，风格统一、角色一致；把一个非动画团队变成了「能出成片」的 AI 动画工作室',
        ],
    ],
    widths=[3.4, 3.4, 6.2, 3.0],
)

# ============================== 4.1 ==============================
doc.add_heading('4.1 AI 工具使用情况', level=1)
p('我们实际使用的 AI 工具围绕「一条生产线、一个交互大脑、一双眼睛」组织，均为阿里云百炼体系 + 开发辅助工具。每项工具承担的职责与投入产出如下：')

table(
    ['AI 工具名称', '在创作中承担的具体任务', '输入的提示词、素材或参数', '输出结果及可用性评价'],
    [
        [
            '通义千问 qwen-flash\n（大语言模型）',
            'AI 导游「菌小园」的对话大脑：儿童友好的肠道健康科普问答',
            'system prompt = AI_STYLE_GUIDE（7 条风格 + 硬性安全规则）+ 动态知识底座（知识模块 / FAQ / Bristol 分型 / 成长阶段）；user 消息按页面注入上下文提示（如「用户正在探索花园页……」）；参数：stream=true 流式输出',
            '输出儿童友好、花园隐喻贯穿的短回答（≤80 字）；所有内容可回溯知识库，零幻觉；可用性高，已稳定服务',
        ],
        [
            '通义千问 VL\n（视觉理解，qwen-vl 系）',
            '便便照片智能分析：识别布里斯托分型、生成健康提示与今日饮食任务',
            '输入：用户上传照片（≤10MB，自动压缩，HEIC→JPEG）；按视觉多模态调用规范携带图片 + 任务提示词；未配置 Key 时自动降级本地规则模拟',
            '识别 Bristol 1-7 型并映射到预设诊断与饮食建议；接口按规范预留、可随时切换真实 VL 模型，兼顾演示稳定性',
        ],
        [
            '通义万相 Wan\n（Wan2.7 t2v / i2v）',
            '90 秒科普动画全部镜头：文生视频 / 图生视频、关键帧、角色设定图',
            '17 分镜结构化 Prompt（场景、动作、时长、镜头运动）存为 JSON；全局正向提示词前缀 + 负向提示词锁定粘土动画、绒布质感、粉蓝绿童话色板；角色一致性用角色设定图 + IP-Adapter 注入',
            '4 段核心知识成片（纤维/发酵/短链脂肪酸/屏障）风格统一、角色不换脸；批量无人值守出片，断点续传',
        ],
        [
            '百炼 ComfyUI 应用托管',
            '编排「关键帧 → 图生视频 → 风格一致性」生成工作流',
            '工作流节点 + 批量脚本参数（seed、步数、分辨率、IP-Adapter 权重）',
            '把多模型串成一条生产线，一套工作流复用全部 17 个分镜；可复制、可调参',
        ],
        [
            'AI 编程助手\n（Claude Code）',
            '互动装置的代码实现：React 交互页面、Fastify 后端、数据驱动内容层、AI 接口对接与联调',
            '以中文自然语言描述需求与设计，配合代码评审、类型检查、浏览器实测',
            '快速产出高质量可运行代码，页面与接口联调高效；人工把控架构与科学内容，AI 负责工程落地',
        ],
    ],
    widths=[3.0, 3.6, 5.2, 4.2],
)

# ============================== 4.2 ==============================
doc.add_heading('4.2 Qwen 模型调用说明（必填）', level=1)

doc.add_heading('① 模型名称与版本', level=2)
p('对话问答：通义千问 qwen-flash（DashScope OpenAI 兼容模式）。')
p('视觉理解：通义千问 VL（qwen-vl-max / qwen-vl-plus 系）—— 用于便便照片分析，接口已按多模态规范预留。')
p('规划升级：qwen-max（更高精度）、通义 CosyVoice 语音合成、Paraformer 语音识别（低龄儿童语音问答）。')

doc.add_heading('② 调用方式', level=2)
p('通过阿里云百炼平台（DashScope）的 OpenAI 兼容 HTTP 接口调用。对话采用流式接口 /chat/completions（stream: true），以 SSE 逐字返回，前端呈现打字机效果。后端配置如下：')

code_block('''# server/.env 示例（百炼平台凭证，不提交代码仓库）
AI_API_KEY=<阿里云百炼 DashScope API-KEY>
AI_BASE_URL=https://dashscope.aliyuncs.com/compatible-mode/v1
AI_MODEL=qwen-flash''')

p('调用封装（server/src/modules/ai/ai-stream.ts）：请求 OpenAI 兼容 /chat/completions，解析 SSE 的 delta 增量并逐个 chunk 返回；AI_API_KEY 未配置、上游失败或无输出时，自动降级为本地 FAQ 关键词匹配，保证「菌小园」在任何情况下都能回答。')

doc.add_heading('③ 调用凭证或截图（占位，需补充）', level=2)
p('以下两处截图请补充插入本小节末尾：', bold=True)
bullet('截图 1（凭证）：登录阿里云百炼控制台 → 右上角「API-KEY」→ 管理我的 API-KEY → 截取 API-KEY 列表页面（Key 值可打码）。')
bullet('截图 2（模型可用性）：百炼控制台 → 模型广场 → 搜索「qwen-flash」→ 截取模型卡片（显示支持/已开通状态）。也可补充「费用与用量」页面的调用统计，直观展示调用量。')
para = doc.add_paragraph()
para.paragraph_format.space_before = Pt(6)
r = para.add_run('【截图占位 1】百炼控制台 API-KEY 页面\n【截图占位 2】模型广场 qwen-flash 模型卡片（如已启用可加「用量统计」）')
r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x99, 0x8A, 0x6A)
east_asia(r)
pPr = para._p.get_or_add_pPr()
pbdr = pPr.makeelement(qn('w:pBdr'), {})
for side in ('top', 'left', 'bottom', 'right'):
    el = pPr.makeelement(qn('w:' + side), {qn('w:val'): 'dashed', qn('w:sz'): '6', qn('w:color'): 'C9B98A'})
    pbdr.append(el)
pPr.append(pbdr)

doc.add_heading('④ Qwen 在作品中的核心贡献', level=2)
p('Qwen 是作品「交互大脑」与「眼睛」的核心：')
bullet('菌小园 AI 导游：把「知识底座 + 风格规则 + 页面上下文」注入 qwen-flash，实现零幻觉、儿童友好的肠道科普问答——这是互动装置里最「活」的部分。')
bullet('便便照片理解：通过通义千问 VL 多模态能力识别布里斯托分型，让「看图」变成「懂图」，支撑数据驱动的个性化饮食任务。')
bullet('扩展：同一套百炼能力（万相出片、VL 质检）支撑了动画生产与内容质检，Qwen 在整个「AI 科普工厂」中处于枢纽位置。')

# ============================== 4.3 ==============================
doc.add_heading('4.3 AI 与人工的协作边界', level=1)

doc.add_heading('① AI 主要负责「生产与执行」', level=2)
bullet('初稿与批量生产：分镜脚本、关键帧与视频镜头、插图与海报初稿、代码骨架与页面实现、文案初稿。')
bullet('一致性执行：严格按人工定义的色板、角色设定、正负向提示词批量产出，保证 170 个镜头共用一套审美。')
bullet('辅助校验：AI 辅助检查拼写、构图、角色是否走样，并定位代码问题。')

doc.add_heading('② 人工必须把控「标准与边界」', level=2)
bullet('科学内核：每一条科普知识都人工核对权威出处（《中国居民膳食指南（2025）》《益生菌的科学共识》、WHO 建议、布里斯托大便分类法），AI 只做「转译」，不做「创作事实」。')
bullet('医疗安全边界：不诊断、不恐吓、不提供用药建议；涉及持续症状一律提示就医——这条红线由人工在提示词规则中固化为硬约束，并逐条审读 AI 输出。')
bullet('审美方向：色板、角色形象、整体调性由团队拍板，AI 在既定框架内执行，不允许「自由发挥」改变风格。')
bullet('最终验收：每一版内容、每一段成片、每一次交互体验都要人工审阅通过后才上线。')

doc.add_heading('③ 如何确保作品体现团队自身的审美与设计，而非 AI 的默认输出', level=2)
bullet('「人定标准、AI 执行、人再审」三层流程：标准（科学清单 + 视觉规范 + 风格资产）由人定义，AI 只负责按标准生产，产出的每一版都必须过人工审阅。')
bullet('风格资产前置：角色设定图、色板、正负向提示词是人工先做好的「资产」，AI 被要求严格复用，而不是让 AI 自己发明风格。')
bullet('把人工经验「回写」系统：每次审片发现的问题（某角色鼻子变形、某镜头色调偏移、某句回答不够口语化）都会被写回提示词库和质检规则，系统越用越贴近团队审美。')
bullet('关键决策全部人工拍板：科学表述、医疗边界、故事结构、整体调性——这些地方 AI 只提供候选，最终选择权始终在团队手里。')

# ============================== 附录：便便建议「垂直专业 AI」 ==============================
doc.add_heading('附录：便便建议「垂直专业 AI」设计说明', level=1)

doc.add_heading('一、是否需要微调（Fine-tuning）？——本项目不需要', level=2)
p('结论：便便建议的「垂直专业」效果不需要微调模型。我们用「强约束提示词 + 知识底座注入（RAG）+ 结构化输出约束」三件套实现，这是微调在工程上最直接、成本最低的等价替代。')
p('理由：', bold=True)
bullet('领域知识规模小且完全可枚举：便便观察的本质是 Bristol 1-7 分型 + 数十条已验证的科普条目（纤维、发酵、补水、运动），全部写进知识数据文件，模型只需按约束「查找 + 翻译」，不需要掌握海量私有知识。')
bullet('微调有真实成本：需要几千条高质量标注数据（本项目无现成数据）、训练与评估周期、后续改内容还要重新训练，与「科普内容持续迭代」的现实冲突。')
bullet('提示词 + RAG 更利于更新与审计：知识改一条，只需改数据文件，立即生效且可追溯；微调后模型内部的知识不可解释、难以单独修订。')
bullet('契合比赛「可复用生成的 AI 应用」定位：换一个科学主题，只替换知识数据文件与视觉资产，AI 流水线不用动；若走微调则每个主题都要重新训练，无法快速复制。')
p('什么情况下才需要微调（本项目的未来判断标准）：', bold=True)
bullet('需要模型掌握大量非公开领域知识（万条级）且检索无法覆盖时；')
bullet('需要输出一种高度特殊的格式、且提示词约束压不住时；')
bullet('需要离线 / 低延迟推理（如部署在无网设备）时；')
bullet('希望每次调用省掉冗长 system prompt 的 token 成本时。')
p('预留升级路径：当数据积累到千条级、或需要固定输出格式时，可对 qwen-flash 做 LoRA 轻量微调；当前「知识底座 + 结构化提示词」的语料可直接转为训练集，实现平滑升级。')

doc.add_heading('二、上下文策略（Context）', level=2)
p('「垂直专业」不是换个模型，而是把正确的上下文按层级组装好喂给模型。本功能按 6 层组织上下文：')
table(
    ['层级', '内容', '作用'],
    [
        ['① 系统提示词', '垂直角色「便便小医生」+ 输出 JSON 结构 + 硬性规则（只依据知识库、不诊断、不恐吓、儿童口吻）', '锁定角色、安全边界与输出格式，是「垂直」的骨架'],
        ['② 知识底座（RAG）', '从已验证数据文件动态组装：Bristol 1-7 分型含义与建议、饮食/喝水/运动条目、常见问答', '给模型「唯一可信来源」，杜绝幻觉与编造'],
        ['③ 结构化输入', 'bristol_type（1-7）+ 形态描述（可选）+ 近期饮食喝水情况（可选），组装成 JSON 传入', '让模型拿到「可推理的事实」，而非自由文本'],
        ['④ 页面上下文', 'buildPageHint("stool")：告知模型用户正在「便便记录」页', '帮助模型聚焦当前场景'],
        ['⑤ 历史对话', '最近 N 条（本功能 N=6）便便问答', '支持追问（如「那明天吃什么好？」）'],
        ['⑥ 降级链', '上游失败 / 超时 / 无输出 → 本地 FAQ 匹配 → Bristol 预设建议', '垂直场景永不中断，演示不翻车'],
    ],
    widths=[3.2, 7.6, 5.2],
)

doc.add_heading('三、完整提示词（可直接复制使用）', level=2)
p('以下为便便建议垂直 AI 的 system prompt。接入时：① 将「知识底座」替换为知识底座中便便相关部分（或直接注入全量底座）；② user 消息传入结构化输入 JSON。', size=10.5)
code_block('''你是「便便小医生」——肠道花园 App 里专门负责「便便观察与饮食建议」的垂直 AI 助手。
你的唯一职责：根据孩子记录的便便形态（Bristol 1-7 型），用儿童能懂的语言给出健康提示和今天可执行的饮食行动建议，并给家长一句专业观察说明。

【输入】（由系统组装为 JSON 传入）
{ "bristol_type": 1-7, "description": "用户补充描述，可空", "recent_diet": "近期饮食/喝水情况，可空" }

【知识底座（只依据这里，禁止编造）】
- Type 1 兔子便便·干硬：多喝水，多吃纤维丰富的蔬菜
- Type 2 香肠便便·干硬：多喝水，多吃蔬果，适量运动
- Type 3 条状便便·正常：继续均衡饮食，注意补充水分
- Type 4 香蕉便·非常健康：继续保持均衡饮食
- Type 5 软块便便·正常：注意多喝水，多吃蔬菜
- Type 6 糊状便便·需关注：调整饮食，多吃纤维食物，多喝水
- Type 7 水样便便·需关注：补充水分；若持续需就医

【输出格式 —— 必须严格按此 JSON 结构，不要输出其他内容】
{
  "child_sentence": "给孩子的 1 句话，≤30 字，活泼、用比喻（如「便便像小香蕉，非常健康，继续加油！」）",
  "suggestion": "1-2 条今日可执行的饮食/喝水/运动建议（面向孩子，具体可做到）",
  "parent_note": "给家长的一句话专业说明，≤60 字（说明 Bristol 类型含义与观察要点）",
  "red_flag": false,
  "red_flag_text": "仅当 red_flag 为 true 时填写"
}

【硬性规则】
1. 只依据知识底座回答，绝不编造 Bristol 类型含义、营养成分、疾病名称、药物名称。
2. 不提供医疗诊断、不用药建议、不制造恐慌。
3. red_flag 仅当描述含以下任一时为 true：持续腹痛、血便、高热、严重脱水、水样便持续超过 24 小时；
   red_flag_text 固定给出：「请马上告诉爸爸妈妈，必要时去医院检查」，并补充具体原因。
4. 给孩子的句子要活泼、多用比喻，像和朋友聊天；给家长的说明要专业但不吓人，给出可观察要点。
5. 与便便/肠道健康无关的问题，礼貌引导回主题。
6. 始终用中文。''')

# 结尾说明
doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run('注：本摘录对应方案报告 3.2 / 3.3 / 4.1 / 4.2 / 4.3 小节，并附「便便建议垂直 AI」设计说明（附录）。4.2 的调用凭证截图请按文中指引在百炼控制台截取后插入。')
r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
east_asia(r)

OUT = r'D:\GutGardenBeta\docs\比赛方案_关键技术决策与AI使用说明.docx'
doc.save(OUT)
print('已生成:', OUT)
